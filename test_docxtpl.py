import os
from src.io.docx_manager import assemble_submission
from docxtpl import DocxTemplate

def test_assemble():
    # create a dummy template to verify it renders without errors
    dummy_template_path = "test_template.docx"

    # We must create a real docx for DocxTemplate to read.
    # The easiest way is to use python-docx to create a blank one, then save it.
    from docx import Document
    d = Document()
    d.add_paragraph("Firm: {{ identity.firm_name }}")
    d.add_paragraph("Has Publishable Clients:")
    d.add_paragraph("{% tr for c in publishable_clients %}")
    d.add_paragraph("Client: {{ c.active_key_client }} - New: {{ c.is_new_client }}")
    d.add_paragraph("{% tr endfor %}")
    d.save(dummy_template_path)

    submission_data = {
        "identity": {"firm_name": "Test Law Firm"},
        "clients": [
            {"active_key_client": "Google", "is_publishable": True, "is_new_client": True},
            {"active_key_client": "Secret Corp", "is_publishable": False, "is_new_client": False}
        ]
    }

    output_dir = "test_output"
    output_path = assemble_submission(dummy_template_path, output_dir, submission_data)

    # Let's read it back
    d2 = Document(output_path)
    text = "\\n".join([p.text for p in d2.paragraphs])
    print("Rendered Text:")
    print(text)

    os.remove(dummy_template_path)
    os.remove(output_path)
    os.rmdir(output_dir)

if __name__ == "__main__":
    test_assemble()
