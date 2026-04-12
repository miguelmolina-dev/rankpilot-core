import os
from src.io.docx_manager import assemble_submission
from docxtpl import DocxTemplate

def test_assemble():
    dummy_template_path = "test_template.docx"
    from docx import Document
    d = Document()
    # To fix the TR error in test we need a real table, so docxtpl recognizes it as a table row loop
    d.add_paragraph("Firm: {{ identity.firm_name }}")
    table = d.add_table(rows=2, cols=2)
    # The {% tr %} tag must only be used inside tables!
    table.rows[0].cells[0].text = "{% tr for c in publishable_clients %}"
    table.rows[0].cells[1].text = "{{ c.active_key_client }}"
    table.rows[1].cells[0].text = "{% tr endfor %}"

    d.save(dummy_template_path)

    submission_data = {
        "identity": {"firm_name": "Test Law Firm"},
        "clients": [
            {"active_key_client": "Google", "is_publishable": True, "is_new_client": True},
            {"active_key_client": "Secret Corp", "is_publishable": False, "is_new_client": False}
        ]
    }

    output_dir = "test_output"
    output_path = assemble_submission(dummy_template_path, output_dir, submission_data)

    d2 = Document(output_path)
    print("Rendered successfully. Number of tables:", len(d2.tables))

    os.remove(dummy_template_path)
    os.remove(output_path)
    os.rmdir(output_dir)

if __name__ == "__main__":
    test_assemble()
