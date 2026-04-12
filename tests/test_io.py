import unittest
import base64
import os
import docx
from fpdf import FPDF
from src.io.base64_handler import decode_base64_document
from src.io.pdf_parser import extract_text_from_pdf
from src.io.docx_manager import extract_text_from_docx

class TestBase64Handler(unittest.TestCase):
    def test_decode_valid_document(self):
        # Create dummy docx content
        dummy_content = b"Dummy docx content"
        b64_string = base64.b64encode(dummy_content).decode('utf-8')

        path = decode_base64_document(b64_string, "test_doc.docx", "/tmp")

        self.assertIsNotNone(path)
        self.assertTrue(os.path.exists(path))

        # Verify content
        with open(path, "rb") as f:
            self.assertEqual(f.read(), dummy_content)

        # Clean up
        os.remove(path)

    def test_decode_invalid_extension(self):
        # Create dummy txt content (not allowed)
        dummy_content = b"Dummy txt content"
        b64_string = base64.b64encode(dummy_content).decode('utf-8')

        path = decode_base64_document(b64_string, "test_doc.txt", "/tmp")

        # Should return None and not create file
        self.assertIsNone(path)

class TestTableExtraction(unittest.TestCase):
    def setUp(self):
        # Create a dummy PDF with a table
        self.pdf_path = "test_table.pdf"
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        pdf.cell(200, 10, text="Normal PDF text", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(50, 10, text="PDFH1", border=1)
        pdf.cell(50, 10, text="PDFH2", border=1, new_x="LMARGIN", new_y="NEXT")
        pdf.cell(50, 10, text="PDFV1", border=1)
        pdf.cell(50, 10, text="PDFV2", border=1, new_x="LMARGIN", new_y="NEXT")
        pdf.output(self.pdf_path)

        # Create a dummy DOCX with a table
        self.docx_path = "test_table.docx"
        doc = docx.Document()
        doc.add_paragraph("Normal DOCX text")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "DOCXH1"
        table.cell(0, 1).text = "DOCXH2"
        table.cell(1, 0).text = "DOCXV1"
        table.cell(1, 1).text = "DOCXV2"
        doc.save(self.docx_path)

    def test_extract_pdf_tables(self):
        text = extract_text_from_pdf(self.pdf_path)
        self.assertIsNotNone(text)
        self.assertIn("Normal PDF text", text)
        self.assertIn("| PDFH1 | PDFH2 |", text)
        self.assertIn("| --- | --- |", text)
        self.assertIn("| PDFV1 | PDFV2 |", text)

    def test_extract_docx_tables(self):
        text = extract_text_from_docx(self.docx_path)
        self.assertIsNotNone(text)
        self.assertIn("Normal DOCX text", text)
        self.assertIn("| DOCXH1 | DOCXH2 |", text)
        self.assertIn("| --- | --- |", text)
        self.assertIn("| DOCXV1 | DOCXV2 |", text)

    def tearDown(self):
        if os.path.exists(self.pdf_path):
            os.remove(self.pdf_path)
        if os.path.exists(self.docx_path):
            os.remove(self.docx_path)
